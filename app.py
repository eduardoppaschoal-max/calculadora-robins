import streamlit as st
from docx import Document
from docx.shared import Pt, RGBColor
from fpdf import FPDF
import io

# --- CONFIGURAÇÃO DA PÁGINA ---
st.set_page_config(
    page_title="ROBINS-I V2 Calculator",
    layout="wide",
    initial_sidebar_state="expanded"
)

# --- FUNÇÕES DE RELATÓRIO (PDF e WORD) ---
def generate_docx(data):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    doc.add_heading(f"Relatório ROBINS-I V2: {data['study_id']}", 0)
    doc.add_paragraph(f"Desfecho: {data['outcome']}")
    doc.add_paragraph(f"Resultado Numérico: {data['numeric_result']}")
    
    # Risco Geral
    doc.add_heading("Julgamento Geral de Risco", level=1)
    p = doc.add_paragraph()
    runner = p.add_run(f"Sugestão do Algoritmo: {data['algo_risk']}")
    runner.bold = True
    
    doc.add_paragraph(f"Decisão Final do Pesquisador: {data['manual_risk']}")
    doc.add_paragraph(f"Justificativa Final: {data['manual_justification']}")

    # Detalhes por Domínio
    doc.add_heading("Detalhamento por Domínio", level=1)
    
    for domain, details in data['domains'].items():
        doc.add_heading(domain, level=2)
        doc.add_paragraph(f"Risco Calculado: {details['risk']}")
        doc.add_paragraph(f"Justificativa do Algoritmo: {details['reason']}")
        doc.add_paragraph("Respostas Selecionadas:")
        for q, a in details['answers'].items():
            doc.add_paragraph(f"  - {q}: {a}", style='List Bullet')

    # Salvar em memória
    bio = io.BytesIO()
    doc.save(bio)
    return bio

def generate_pdf(data):
    class PDF(FPDF):
        def header(self):
            self.set_font('Arial', 'B', 15)
            self.cell(0, 10, f"Relatorio ROBINS-I V2: {data['study_id']}", 0, 1, 'C')
            self.ln(10)

    pdf = PDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)

    def clean_text(text):
        return str(text).encode('latin-1', 'replace').decode('latin-1')

    # Cabeçalho Info
    pdf.set_font("Arial", 'B', 12)
    pdf.cell(0, 10, clean_text(f"Desfecho: {data['outcome']}"), 0, 1)
    pdf.cell(0, 10, clean_text(f"Resultado: {data['numeric_result']}"), 0, 1)
    pdf.ln(5)

    # Risco Geral
    pdf.set_font("Arial", 'B', 14)
    pdf.cell(0, 10, clean_text("Julgamento Geral"), 0, 1)
    pdf.set_font("Arial", '', 12)
    pdf.multi_cell(0, 10, clean_text(f"Algoritmo: {data['algo_risk']}"))
    pdf.multi_cell(0, 10, clean_text(f"Decisao Pesquisador: {data['manual_risk']}"))
    pdf.multi_cell(0, 10, clean_text(f"Justificativa: {data['manual_justification']}"))
    pdf.ln(5)

    # Domínios
    pdf.set_font("Arial", 'B', 14)
    pdf.cell(0, 10, clean_text("Detalhamento por Dominio"), 0, 1)
    
    pdf.set_font("Arial", '', 11)
    for domain, details in data['domains'].items():
        pdf.set_font("Arial", 'B', 12)
        pdf.cell(0, 10, clean_text(domain), 0, 1)
        pdf.set_font("Arial", '', 11)
        pdf.cell(0, 8, clean_text(f"Risco: {details['risk']}"), 0, 1)
        pdf.multi_cell(0, 8, clean_text(f"Motivo: {details['reason']}"))
        pdf.ln(2)

    return pdf.output(dest="S").encode("latin-1")

# --- FUNÇÕES AUXILIARES DE UI ---
def get_risk_color(risk, domain_name=""):
    r = str(risk).upper()
    d = str(domain_name).upper()
    
    # 1. Checagem de Baixo Risco
    if "LOW" in r or "BAIXO RISCO" in r:
        # REGRA ESPECIAL: Domínio 1 é sempre Amarelo (exceto preocupações)
        if "DOMÍNIO 1" in d:
            return "#D4AC0D"  # Amarelo Escuro
        # REGRA PADRÃO: Outros domínios (2, 3, etc) são Verdes
        return "#27AE60"      # Verde Esmeralda
        
    # 2. Outros Níveis de Risco
    elif "MODERATE" in r or "MODERADO" in r: 
        return "#E67E22"  # Laranja
    elif "SERIOUS" in r or "SÉRIO" in r or "SERIO" in r: 
        return "#C0392B"  # Vermelho
    elif "CRITICAL" in r or "CRÍTICO" in r or "CRITICO" in r: 
        return "#000000"  # Preto
        
    # 3. Padrão (Pendente ou erro)
    return "gray"

def display_risk_card(domain, risk, justification):
    # O SEGREDO ESTÁ AQUI: Passamos 'domain' para get_risk_color saber se aplica a regra do Amarelo ou Verde
    color = get_risk_color(risk, domain_name=domain)
    
    st.markdown(f"""
    <div style="padding: 10px; border-left: 5px solid {color}; background-color: #f0f2f6; margin-bottom: 10px;">
        <strong>{domain}:</strong> <span style="color: {color}; font-weight: bold;">{risk}</span><br>
        <em style="font-size: 0.9em;">{justification}</em>
    </div>
    """, unsafe_allow_html=True)

# --- BARRA LATERAL ---
with st.sidebar:
    st.header("Dados do Estudo")
    study_id = st.text_input("ID do Estudo / Autor", value="Estudo Exemplo")
    outcome = st.text_input("Desfecho Avaliado", value="Mortalidade")
    numeric_result = st.text_input("Resultado Numérico", value="RR 1.5")
    st.divider()
    st.info("Ferramenta baseada no ROBINS-I V2 (Nov 2025).")

st.title("ROBINS-I V2: Calculadora de Risco de Viés")
if study_id:
    st.subheader(f"Avaliando: {study_id}")

# --- 1. TRIAGEM E CONTEXTO ---
st.header("1. Considerações Preliminares (Triagem)")
col_b1, col_b2, col_b3 = st.columns(3)
with col_b1: b1 = st.selectbox("B1. Os autores fizeram alguma tentativa de controlar fatores de confusão no resultado avaliado?", ["Selecione...", "Y", "PY", "PN", "N"])
with col_b2: b2 = st.selectbox("B2. Se N/PN para B1: Existe potencial suficiente para fatores de confusão que impeçam a consideração deste resultado posteriormente?", ["Selecione...", "N", "PN", "Y", "PY"])
with col_b3: b3 = st.selectbox("B3. O método de medição do resultado foi inadequado?", ["Selecione...", "N", "PN", "Y", "PY"])

# TRAVA DE SEGURANÇA
if b2 in ["Y", "PY"] or b3 in ["Y", "PY"]:
    st.error("🚨 RISCO CRÍTICO DETECTADO NA TRIAGEM (B2 ou B3). Pare a avaliação aqui.")
    st.stop()
st.divider()

# SELEÇÃO DE VARIANTE (C4)
st.markdown("### Contexto da Análise")
c4 = st.radio(
    "C4. A análise levou em consideração as mudanças entre as estratégias de intervenção comparadas durante o acompanhamento, ou outros desvios de protocolo durante o acompanhamento?", 
    ["Não (Intention-to-treat / Atribuição)", "Sim (Per-protocol / Adesão)"]
)
is_variant_a = "Não" in c4

# Inicialização de variáveis globais
report_data = {
    "study_id": study_id,
    "outcome": outcome,
    "numeric_result": numeric_result,
    "domains": {}
}
risks = {}
reasons = {}

# --- DOMÍNIO 1: CONFUSÃO ---
st.header("Domínio 1: Viés devido a Confusão")

if is_variant_a:
    st.caption("Variante A (Intention-to-treat): Foco na atribuição da intervenção.")
    c1, c2 = st.columns(2)

    # COLUNA 1
    with c1:
        help_1_1 = """
        CONTEXTO: Fatores da avaliação preliminar.
        - Y / PY: Todos fatores importantes foram controlados adequadamente.
        - WN (Não, não substancial): A maioria foi controlada. Viés residual provável é pequeno.
        - SN (Não, substancial): Fator importante NÃO controlado com provável impacto no resultado.
        """
        q1_1 = st.selectbox(
            "1.1 Os autores controlaram todos os importantes fatores de confusão que isso se mostrou necessário?", 
            ["Selecione...", "Y", "PY", "WN", "SN", "NI"], 
            help=help_1_1
        )
        
        # 1.4 SEMPRE visível
        help_1_4 = """
        CONTEXTO: Controles Negativos.
        - Y / PY (Alerta): Controle negativo mostrou associação (viés).
        - N / PN (Neutro): Sem problemas detectados.
        - NA: Não foram usados controles negativos.
        """
        q1_4 = st.selectbox(
            "1.4 O uso de controles negativos sugeriu a presença de fatores de confusão não controlados?", 
            ["Selecione...", "Y", "PY", "N", "PN", "NA"],
            help=help_1_4
        )

    # COLUNA 2
    with c2:
        # Visibilidade dinâmica: 1.2 e 1.3 só aparecem se houve tentativa de controle
        enable_details = q1_1 in ["Y", "PY", "WN"]
        
        if enable_details:
            help_1_2 = """
            CONTEXTO: Validade das medidas usadas.
            - Y / PY: Medidas válidas/confiáveis usadas.
            - WN / SN: Medidas com problemas de validade ou confiabilidade.
            - NA: Se não havia fatores de confusão.
            """
            q1_2 = st.selectbox(
                "1.2 Os fatores de confusão que foram controlados foram medidos de forma válida e confiável?", 
                ["Selecione...", "Y", "PY", "WN", "SN", "NI", "NA"],
                help=help_1_2
            )
            
            help_1_3 = """
            CONTEXTO: Ajuste Excessivo (Over-adjustment).
            - Y / PY (Risco): Controlaram mediadores ou colisores.
            - N / PN (Ideal): Não controlaram variáveis indevidas.
            """
            q1_3 = st.selectbox(
                "1.3 Os autores controlaram alguma variável pós-intervenção que poderia ter sido afetada pela intervenção?", 
                ["Selecione...", "Y", "PY", "N", "PN", "NI", "NA"],
                help=help_1_3
            )
        else:
            q1_2 = "NA"
            q1_3 = "NA"

    d1_risk = "PENDENTE"
    d1_reason = "Aguardando respostas..."
    
    # --- ALGORITMO OTIMIZADO DOMÍNIO 1 (Early Exit) ---
    # Prioridade para riscos CRÍTICOS e SÉRIOS sem exigir preenchimento total se não necessário.

    # 1. ATALHO CRÍTICO A: Falha Controle (SN/NI) + Viés Confirmado (1.4 Y/PY)
    if (q1_1 in ["SN", "NI"]) and (q1_4 in ["Y", "PY"]):
        d1_risk = "CRITICAL"
        d1_reason = "Determinante: Falha no controle (1.1) confirmada por controles negativos (1.4)."

    # 2. ATALHO CRÍTICO B: Ajuste Excessivo (1.3 Y/PY) + Viés Confirmado (1.4 Y/PY)
    elif (q1_1 in ["Y", "PY", "WN"]) and (q1_3 in ["Y", "PY"]) and (q1_4 in ["Y", "PY"]):
        d1_risk = "CRITICAL"
        d1_reason = "Determinante: Ajuste excessivo (1.3) confirmado por controles negativos (1.4)."

    # 3. ATALHO SÉRIO: Erro de Medição Grave (Sem Ajuste Excessivo)
    elif (q1_1 in ["Y", "PY", "WN"]) and (q1_3 in ["N", "PN", "NI", "NA"]) and (q1_2 in ["SN", "NI"]):
        d1_risk = "SERIOUS"
        d1_reason = "Determinante: Erro substancial na medição dos fatores (1.2)."

    # 4. CÁLCULO DETALHADO (Se não caiu nos atalhos)
    else:
        can_calculate = False
        
        # Se Falha Controle: Precisa de 1.4
        if q1_1 in ["SN", "NI"] and q1_4 != "Selecione...":
            can_calculate = True
            
        # Se Controle OK: Precisa de 1.2, 1.3 e 1.4
        elif q1_1 in ["Y", "PY", "WN"] and (q1_2 != "Selecione...") and (q1_3 != "Selecione...") and (q1_4 != "Selecione..."):
            can_calculate = True

        if can_calculate:
            # CAMINHO A: FALHA NO CONTROLE (1.1 = SN/NI)
            if q1_1 in ["SN", "NI"]:
                # Se não caiu no Atalho Crítico A, 1.4 é N/PN/NA -> Sério
                d1_risk = "SERIOUS"
                d1_reason = "Falha substancial no controle (1.1). Controles negativos não agravaram para crítico."

            # CAMINHO B: CONTROLE TENTADO (1.1 = Y/PY/WN)
            else:
                is_critical = False
                is_serious = False

                # --- ANÁLISE DE AJUSTE EXCESSIVO (1.3 = Y/PY) ---
                if q1_3 in ["Y", "PY"]:
                    # Já testamos 1.4=Y/PY no Atalho Crítico B.
                    # Resta testar Medição Ruim.
                    if q1_2 in ["SN", "WN", "NI"]:
                        d1_risk = "CRITICAL"
                        d1_reason = "Ajuste excessivo (1.3) agravado por medição insuficiente (1.2)."
                        is_critical = True
                    else:
                        d1_risk = "SERIOUS"
                        d1_reason = "Ajuste excessivo de variáveis (1.3), mitigado por boa medição."
                        is_serious = True
                
                # --- SEM AJUSTE EXCESSIVO (1.3 = N/PN/NA) ---
                else:
                    # Controles Negativos Apitando
                    if q1_4 in ["Y", "PY"]:
                        d1_risk, d1_reason = "SERIOUS", "Controles negativos sugerem viés, apesar do bom controle inicial."
                        is_serious = True
                    
                    # Erro de Medição Grave (Já tratado no Atalho 3, mas reforço lógica aqui)
                    elif q1_2 in ["SN", "NI"]:
                        d1_risk, d1_reason = "SERIOUS", "Erro substancial na medição dos fatores (1.2)."
                        is_serious = True
                
                if not is_critical and not is_serious:
                    if q1_2 == "WN" or q1_1 == "WN":
                        d1_risk = "MODERATE"
                        d1_reason = "Preocupações menores com confusão residual ou erro de medição."
                    else:
                        d1_risk = "LOW"
                        d1_reason = "Baixo risco de viés devido a confusão."

    # --- TRADUÇÃO DO RESULTADO (SÓ PARA DOMÍNIO 1) ---
    if d1_risk == "LOW":
        d1_risk = "Baixo risco, exceto por preocupações com confusão"

    risks["D1"] = d1_risk
    reasons["D1"] = d1_reason
    
    # Observe que aqui listamos EXATAMENTE as variáveis da Variante A
    report_data["domains"]["Domínio 1"] = {
        "risk": d1_risk, 
        "reason": d1_reason, 
        "answers": {"1.1": q1_1, "1.2": q1_2, "1.3": q1_3, "1.4": q1_4}
    }
    
    display_risk_card("Domínio 1", d1_risk, d1_reason)

else:
    # --- VARIANTE B (Quando C4 = Sim / Per-protocol) ---
    st.caption("Variante B (Efeito da adesão à intervenção): Foco em confusão variável no tempo.")
    
    c1, c2 = st.columns(2)

    with c1:
        # PERGUNTA 1.1
        help_1_1 = """
        Métodos apropriados para controlar fatores de confusão variáveis no tempo ('métodos g') incluem aqueles baseados na ponderação por probabilidade inversa. 
        Modelos de regressão padrão que incluem fatores de confusão variáveis no tempo podem ser problemáticos quando esses fatores são afetados por intervenções anteriores.
        """
        q1_1 = st.selectbox(
            "1.1 Os autores utilizaram um método de análise apropriado para controlar os fatores de confusão variáveis ao longo do tempo, bem como os fatores de confusão basais?", 
            ["Selecione...", "Y", "PY", "PN", "N", "NI"], 
            help=help_1_1
        )

        # PERGUNTA 1.5 (Sempre visível, pois é crucial para a maioria dos caminhos)
        help_1_5 = """
        A utilização de um "controle negativo" pode sugerir fatores de confusão não controlados.
        - N: Não houve sinal de viés (ou não foi feito).
        - Y/PY: Controles negativos indicaram viés.
        """
        q1_5 = st.selectbox(
            "1.5 O uso de controles negativos, ou outras considerações, sugeriu a presença de fatores de confusão não controlados significativos?", 
            ["Selecione...", "Y", "PY", "PN", "N"], 
            help=help_1_5
        )

    with c2:
        # VISIBILIDADE DINÂMICA
        q1_2 = "NA"
        q1_3 = "NA"
        q1_4 = "NA"

        # Caminho Método Adequado (Y/PY)
        if q1_1 in ["Y", "PY"]:
            help_1_2 = """
            - Y/PY: Todos fatores importantes (basais e variáveis no tempo) controlados.
            - WN: Maioria controlada, viés residual provável é pequeno.
            - SN: Fator importante não controlado.
            """
            q1_2 = st.selectbox(
                "1.2 Os autores controlaram todos os importantes fatores de confusão basais e variáveis ao longo do tempo para os quais isso era necessário?",
                ["Selecione...", "NA", "Y", "PY", "WN", "SN", "NI"],
                help=help_1_2
            )

            # 1.3 só aparece se 1.2 não foi uma falha total
            if q1_2 in ["Y", "PY", "WN"]:
                q1_3 = st.selectbox(
                    "1.3 Os fatores de confusão que foram controlados foram medidos de forma válida e confiável?",
                    ["Selecione...", "NA", "Y", "PY", "WN", "SN", "NI"],
                    help="Se a validade/confiabilidade não for citada, avalie a subjetividade."
                )
        
        # Caminho Método Inadequado (N/PN/NI)
        elif q1_1 in ["N", "PN", "NI"]:
            help_1_4 = """
            Verificação de Viés de Colisor.
            - Y/PY: Controlaram variáveis pós-intervenção em método padrão (CRÍTICO).
            - N/PN: Não controlaram (Sério, mas evita colisor).
            """
            q1_4 = st.selectbox(
                "1.4 Os autores controlaram fatores que variam ao longo do tempo ou outras variáveis medidas após o início da intervenção?",
                ["Selecione...", "NA", "Y", "PY", "PN", "N", "NI"],
                help=help_1_4
            )

    d1_risk = "PENDENTE"
    d1_reason = "Aguardando respostas..."

    # --- ALGORITMO INTELIGENTE (VARIANTE B) ---
    
    # 1. ATALHO DE RISCO CRÍTICO (Independente de 1.5)
    # Viés de Colisor: Método Ruim + Controle de Pós-intervenção
    if q1_1 in ["N", "PN", "NI"] and q1_4 in ["Y", "PY"]:
        d1_risk = "CRITICAL"
        d1_reason = "Determinante: Método inadequado com ajuste por variáveis pós-intervenção (Viés de Colisor)."
    
    # 2. CÁLCULO PARA OS DEMAIS CASOS (Requer 1.5 preenchido)
    elif q1_5 != "Selecione...":
        
        # --- CAMINHO A: MÉTODO INADEQUADO (1.1 N/PN/NI) ---
        if q1_1 in ["N", "PN", "NI"]:
            # Se chegou aqui, 1.4 não é Y/PY (pois cairia no atalho acima)
            if q1_4 != "Selecione...":
                if q1_5 in ["Y", "PY"]:
                    d1_risk = "CRITICAL"
                    d1_reason = "Método inadequado e controles negativos indicam confusão não controlada."
                else:
                    d1_risk = "SERIOUS"
                    d1_reason = "Método de análise inadequado para adesão (falha em ajustar confusão variável no tempo)."

        # --- CAMINHO B: MÉTODO ADEQUADO (1.1 Y/PY) ---
        elif q1_1 in ["Y", "PY"]:
            
            # Precisamos verificar se temos dados suficientes de 1.2 e 1.3
            can_calc_b = False
            if q1_2 in ["SN", "NI"]: can_calc_b = True # Falha controle já define
            elif q1_2 in ["Y", "PY", "WN"] and q1_3 != "Selecione...": can_calc_b = True

            if can_calc_b:
                is_critical = False
                is_serious = False

                # 1. Checagem de CRÍTICO (Falhas Graves + Viés Confirmado)
                if q1_5 in ["Y", "PY"]:
                    if q1_2 in ["SN", "NI"]:
                        d1_risk, d1_reason = "CRITICAL", "Falha substancial no controle confirmada por controles negativos."
                        is_critical = True
                    elif q1_3 in ["SN", "NI"]:
                        d1_risk, d1_reason = "CRITICAL", "Medição inválida confirmada por viés em controles negativos."
                        is_critical = True

                if not is_critical:
                    # 2. Checagem de SÉRIO
                    
                    # Grupo B: Falhas Substanciais (Sem confirmação externa de viés)
                    if q1_2 in ["SN", "NI"]:
                        d1_risk, d1_reason = "SERIOUS", "Falha substancial no controle de fatores de confusão."
                        is_serious = True
                    elif q1_3 in ["SN", "NI"]:
                        d1_risk, d1_reason = "SERIOUS", "Falha substancial na medição dos fatores de confusão."
                        is_serious = True
                    
                    # Grupo C: Viés Confirmado (Agravante para estudos Bons/Moderados)
                    elif q1_5 in ["Y", "PY"]:
                        d1_risk = "SERIOUS"
                        d1_reason = "Controles negativos sugerem viés, apesar da metodologia aparentemente adequada."
                        is_serious = True
                    
                    if not is_serious:
                        # 3. MODERADO
                        # Ressalvas em Controle (1.2 WN) ou Medição (1.3 WN)
                        if q1_2 == "WN" or q1_3 == "WN":
                            d1_risk = "MODERATE"
                            d1_reason = "Ressalvas menores no controle ou medição dos fatores de confusão."
                        
                        # 4. BAIXO
                        else:
                            d1_risk = "LOW"
                            d1_reason = "Baixo risco de viés (G-methods aplicados corretamente)."

# --- AJUSTE DE TEXTO (TRADUÇÃO) ---
    # Garante que o texto exibido seja o padrão do ROBINS-I para Domínio 1
    if d1_risk == "LOW":
        d1_risk = "Baixo risco, exceto por preocupações com confusão"

    # Salva nos dados globais
    risks["D1"] = d1_risk
    reasons["D1"] = d1_reason
    
    # CORREÇÃO AQUI: Listamos explicitamente todas as variáveis da Variante B
    report_data["domains"]["Domínio 1"] = {
        "risk": d1_risk, 
        "reason": d1_reason, 
        "answers": {"1.1": q1_1, "1.2": q1_2, "1.3": q1_3, "1.4": q1_4, "1.5": q1_5}
    }
    
    display_risk_card("Domínio 1", d1_risk, d1_reason)

st.divider()

# --- DOMÍNIO 2: CLASSIFICAÇÃO ---
st.header("Domínio 2: Viés na Classificação das Intervenções")

# Layout: 2.1 (Tempo Imortal) e condicionais na esquerda; 2.4 e 2.5 (Influência/Erro) na direita.
c1_d2, c2_d2 = st.columns(2)

with c1_d2:
    # --- BLOCO TEMPO IMORTAL (2.1, 2.2, 2.3) ---
    st.markdown("###### Definição da Intervenção")
    
    # 2.1 (Sempre visível)
    help_2_1 = """
    No ensaio alvo, o acompanhamento começa na elegibilidade. Em estudos não randomizados, algumas estratégias não são distinguíveis no início (ex: "operar em 6 meses" vs "esperar"). 
    Classificar participantes baseando-se em eventos futuros gera "viés de tempo imortal".
    """
    q2_1 = st.selectbox(
        "2.1 As estratégias de intervenção eram distinguíveis no momento em que o acompanhamento teria começado?", 
        ["Selecione...", "Y", "PY", "PN", "N", "NI"],
        help=help_2_1
    )

    # Lógica de Visibilidade em Cascata (2.2 e 2.3)
    q2_2 = "NA"
    q2_3 = "NA"

    # 2.2 só aparece se 2.1 for problemático
    if q2_1 in ["N", "PN", "NI"]:
        help_2_2 = """
        Se o período de indistinção for curto em relação ao acompanhamento total, poucos eventos ocorrerão nele, limitando o risco de viés.
        """
        q2_2 = st.selectbox(
            "2.2 Todos ou quase todos os eventos ocorreram após a intervenção ser distinguível?",
            ["Selecione...", "Y", "PY", "PN", "N", "NI"],
            help=help_2_2
        )
        
        # 2.3 só aparece se 2.2 TAMBÉM for problemático
        if q2_2 in ["N", "PN", "NI"]:
            help_2_3 = """
            Métodos estatísticos avançados (ponderação por censura clonal, g-formula) podem corrigir problemas de estratégias indistinguíveis.
            - SY: Sim, totalmente.
            - WY: Sim, parcialmente.
            """
            q2_3 = st.selectbox(
                "2.3 A análise evitou problemas decorrentes de estratégias indistinguíveis?",
                ["Selecione...", "SY", "WY", "PN", "N", "NI"],
                help=help_2_3
            )

with c2_d2:
    # --- BLOCO CLASSIFICAÇÃO (2.4, 2.5) - SEMPRE VISÍVEIS ---
    st.markdown("###### Validade da Classificação")

    # 2.4 (Sempre visível)
    help_2_4 = """
    A classificação da intervenção foi influenciada pelo conhecimento do desfecho?
    (Comum em estudos retrospectivos onde o avaliador sabe quem morreu/sobreviveu ao classificar o tratamento).
    - SY: Sim, totalmente (Risco Alto).
    - WY: Sim, parcialmente.
    """
    q2_4 = st.selectbox(
        "2.4 A classificação da intervenção foi influenciada pelo conhecimento do desfecho?", 
        ["Selecione...", "SY", "WY", "PN", "N", "NI"], 
        help=help_2_4
    )

    # 2.5 (Sempre visível)
    help_2_5 = """
    Houve erros na classificação do status da intervenção?
    (Critérios ambíguos ou registros incompletos. Se o erro for aleatório, tende a viés para o nulo).
    """
    q2_5 = st.selectbox(
        "2.5 Houve erros na classificação do status da intervenção?",
        ["Selecione...", "Y", "PY", "PN", "N", "NI"],
        help=help_2_5
    )

d2_risk = "PENDENTE"
d2_reason = "Aguardando respostas..."

# --- ALGORITMO INTELIGENTE DOMÍNIO 2 ---

# Passo 1: Determinar o "Contexto de Entrada" (Status do Tempo Imortal)
# SAFE: Problema resolvido ou inexistente.
# PARTIAL: Problema parcialmente resolvido (2.3 WY/NI).
# BAD: Problema não resolvido (2.3 N/PN).

entry_context = "PENDING"

if q2_1 in ["Y", "PY"]: entry_context = "SAFE"
elif q2_1 in ["N", "PN", "NI"]:
    if q2_2 in ["Y", "PY"]: entry_context = "SAFE"
    elif q2_2 in ["N", "PN", "NI"]:
        if q2_3 == "SY": entry_context = "SAFE"
        elif q2_3 in ["WY", "NI"]: entry_context = "PARTIAL"
        elif q2_3 in ["N", "PN"]: entry_context = "BAD"
        # Se 2.3 for Selecione..., continua PENDING

# Se 2.4 ou 2.5 não foram respondidos, marcamos como pendente para cálculo final,
# MAS tentaremos calcular riscos críticos imediatos abaixo.
inputs_missing = (q2_4 == "Selecione...") or (q2_5 == "Selecione...")

# Passo 2: Cálculo de Risco
# A lógica tenta encontrar o pior cenário possível com os dados disponíveis.

calculated = False

# --- VERIFICAÇÃO DE RISCO CRÍTICO (Prioridade Máxima) ---
# 1. Influência Total do Desfecho + Erro de Classificação (Independe da Entrada)
if q2_4 == "SY" and q2_5 in ["Y", "PY", "NI"]:
    d2_risk, d2_reason = "CRITICAL", "Determinante: Classificação totalmente influenciada pelo desfecho com erros adicionais."
    calculated = True

# 2. Entrada Ruim/Parcial + Influência do Desfecho (Independe de 2.5)
elif entry_context in ["BAD", "PARTIAL"] and q2_4 in ["SY", "WY", "NI"]:
    d2_risk, d2_reason = "CRITICAL", "Determinante: Problema de tempo imortal não resolvido somado à influência do desfecho."
    calculated = True

# 3. Entrada Ruim + Erro de Classificação (Se 2.4 for ok ou pendente)
elif entry_context == "BAD" and q2_5 in ["Y", "PY", "NI"]:
    d2_risk, d2_reason = "CRITICAL", "Determinante: Problema de tempo imortal não resolvido com erros de classificação."
    calculated = True


if not calculated and not inputs_missing:
    # --- VERIFICAÇÃO DE RISCO SÉRIO ---
    is_serious = False
    
    # 4. Entrada Segura + Incerteza Desfecho + Erro Classificação
    if entry_context == "SAFE" and q2_4 in ["WY", "NI"] and q2_5 in ["Y", "PY", "NI"]:
        d2_risk, d2_reason = "SERIOUS", "Combinação de possível influência do desfecho e erros de classificação."
        is_serious = True
        
    # 5. Entrada Segura + Influência Total (Sem erro 2.5)
    elif entry_context == "SAFE" and q2_4 == "SY":
        d2_risk, d2_reason = "SERIOUS", "Classificação influenciada pelo desfecho (viés diferencial)."
        is_serious = True
        
    # 6. Entrada Parcial + Erro de Classificação
    elif entry_context == "PARTIAL" and q2_5 in ["Y", "PY", "NI"]:
        d2_risk, d2_reason = "SERIOUS", "Correção apenas parcial do tempo imortal com erros de classificação."
        is_serious = True
        
    # 7. Entrada Ruim (Pura)
    elif entry_context == "BAD":
        d2_risk, d2_reason = "SERIOUS", "Problema de tempo imortal (intervenções indistinguíveis) não corrigido."
        is_serious = True

    if not is_serious:
        # --- VERIFICAÇÃO DE RISCO MODERADO ---
        is_moderate = False
        
        # 8. Entrada Segura + Erro de Classificação (Puro)
        if entry_context == "SAFE" and q2_5 in ["Y", "PY", "NI"]:
            d2_risk, d2_reason = "MODERATE", "Erros de classificação não-diferenciais (provável viés para o nulo)."
            is_moderate = True
            
        # 9. Entrada Segura + Incerteza Influência
        elif entry_context == "SAFE" and q2_4 in ["WY", "NI"]:
            d2_risk, d2_reason = "MODERATE", "Dúvida leve sobre influência do desfecho."
            is_moderate = True
            
        # 10. Entrada Parcial (Pura)
        elif entry_context == "PARTIAL":
             d2_risk, d2_reason = "MODERATE", "Correção do tempo imortal foi apenas parcial (WY/NI em 2.3)."
             is_moderate = True
        
        if not is_moderate:
            # --- BAIXO RISCO ---
            if entry_context == "SAFE" and q2_4 in ["N", "PN"] and q2_5 in ["N", "PN"]:
                d2_risk, d2_reason = "LOW", "Intervenção bem definida e classificada sem viés."
            else:
                # Fallback caso a lógica de entrada falhe (ex: entry_context ainda PENDING)
                d2_risk = "PENDENTE"

risks["D2"] = d2_risk
reasons["D2"] = d2_reason

report_data["domains"]["Domínio 2"] = {
    "risk": d2_risk, 
    "reason": d2_reason, 
    "answers": {"2.1": q2_1, "2.2": q2_2, "2.3": q2_3, "2.4": q2_4, "2.5": q2_5}
}
display_risk_card("Domínio 2", d2_risk, d2_reason)

st.divider()

# --- DOMÍNIO 3: SELEÇÃO DOS PARTICIPANTES ---
st.header("Domínio 3: Viés devido à Seleção dos Participantes")

st.markdown("""
Este domínio avalia se a exclusão de participantes ou o tempo de acompanhamento introduz viés. 
O Bloco C (Correção) só será ativado se forem detectados problemas sérios nas partes A ou B.
""")

# Layout Bipartido: Coluna A (Início) e Coluna B (Pós-Início)
c1_d3, c2_d3 = st.columns(2)

# --- PARTE A: Início do Acompanhamento ---
with c1_d3:
    st.subheader("A. Início do Acompanhamento")
    
    help_3_1 = """
    O acompanhamento coincidiu com o início da intervenção?
    - Y/PY: Sim (Ideal).
    - WN: Não, lacuna irrelevante.
    - SY: Início muito tardio (Risco Sério).
    """
    q3_1 = st.selectbox(
        "3.1 Os participantes foram acompanhados desde o início da intervenção?",
        ["Selecione...", "Y", "PY", "WN", "SY", "NI"],
        help=help_3_1
    )

    q3_2 = "NA"
    if q3_1 in ["Y", "PY"]:
        help_3_2 = """
        Eventos precoces foram excluídos?
        - N/PN: Não (Bom).
        - Y/PY: Sim (Risco Moderado).
        """
        q3_2 = st.selectbox(
            "3.2 Os eventos de desfecho precoces foram excluídos da análise?",
            ["Selecione...", "Y", "PY", "PN", "N", "NI"],
            help=help_3_2
        )

# --- PARTE B: Seleção Pós-Início ---
with c2_d3:
    st.subheader("B. Seleção Pós-Início")
    
    help_3_3 = """
    A inclusão foi baseada em características medidas APÓS o início da intervenção?
    - N/PN: Não (Ideal).
    - Y/PY: Sim (Potencial Viés).
    """
    q3_3 = st.selectbox(
        "3.3 A seleção foi baseada em características pós-intervenção?",
        ["Selecione...", "Y", "PY", "PN", "N", "NI"],
        help=help_3_3
    )

    q3_4 = "NA"
    q3_5 = "NA"
    
    if q3_3 in ["Y", "PY"]:
        help_3_4 = """
        Essas características estão associadas à intervenção?
        - N/PN: Não (Risco Baixo).
        - Y/PY: Sim.
        - NI: Sem informação (Risco Moderado).
        """
        q3_4 = st.selectbox(
            "3.4 As variáveis de seleção estão associadas à intervenção?",
            ["Selecione...", "Y", "PY", "PN", "N", "NI"],
            help=help_3_4
        )
        
        if q3_4 in ["Y", "PY", "NI"]:
            help_3_5 = """
            Essas variáveis são influenciadas pelo desfecho?
            - Y/PY: Sim (Risco Sério).
            - N/PN/NI: Não ou Sem Info (Risco Moderado).
            """
            q3_5 = st.selectbox(
                "3.5 As variáveis de seleção são influenciadas pelo desfecho?",
                ["Selecione...", "Y", "PY", "PN", "N", "NI"],
                help=help_3_5
            )

# --- CÁLCULO PROVISÓRIO (Para decidir se mostra o Bloco C) ---
temp_risk_a = "PENDING"
temp_risk_b = "PENDING"

# Lógica Risco A
if q3_1 == "SY": temp_risk_a = "SERIOUS"
elif q3_1 in ["WN", "NI"]: temp_risk_a = "MODERATE"
elif q3_1 in ["Y", "PY"]:
    if q3_2 in ["Y", "PY"]: temp_risk_a = "MODERATE"
    elif q3_2 in ["N", "PN", "NI"]: temp_risk_a = "LOW"

# Lógica Risco B
if q3_3 in ["N", "PN", "NI"]: temp_risk_b = "LOW"
elif q3_3 in ["Y", "PY"]:
    if q3_4 in ["N", "PN"]: temp_risk_b = "LOW"
    elif q3_4 in ["NI"]: temp_risk_b = "MODERATE"
    elif q3_4 in ["Y", "PY"]:
        if q3_5 in ["Y", "PY"]: temp_risk_b = "SERIOUS"
        elif q3_5 in ["N", "PN", "NI"]: temp_risk_b = "MODERATE"

# Combinação Provisória
is_provisional_serious = (temp_risk_a == "SERIOUS") or (temp_risk_b == "SERIOUS")

# --- BLOCO C: CORREÇÃO (Condicional) ---
q3_6 = "NA"
q3_7 = "NA"
q3_8 = "NA"

if is_provisional_serious:
    st.divider()
    st.markdown("###### C. Análise e Correção (Ativado: Risco Sério Detectado)")
    st.caption("Problemas sérios identificados. Responda abaixo para verificar correção.")

    help_3_6 = "A análise usou métodos (ex: IPW, ajuste) para corrigir o viés de seleção?"
    q3_6 = st.selectbox(
        "3.6 A análise corrigiu o viés de seleção?",
        ["Selecione...", "Y", "PY", "PN", "N", "NI"],
        help=help_3_6
    )

    if q3_6 in ["N", "PN", "NI"]:
        q3_7 = st.selectbox(
            "3.7 Análises de sensibilidade demonstram impacto mínimo do viés?",
            ["Selecione...", "Y", "PY", "PN", "N", "NI"],
            help="Se Sim (Y/PY), o risco cai para Moderado."
        )
        
        if q3_7 in ["N", "PN", "NI"]:
            q3_8 = st.selectbox(
                "3.8 O viés de seleção é provável de ser severo?",
                ["Selecione...", "Y", "PY", "PN", "N", "NI"],
                help="Se Sim (Y/PY), o risco se torna CRÍTICO."
            )

# --- ALGORITMO FINAL DOMÍNIO 3 ---
d3_risk = "PENDENTE"
d3_reason = "Aguardando respostas..."

# Verifica se o fluxo foi completado
flow_complete = False
if temp_risk_a != "PENDING" and temp_risk_b != "PENDING":
    if not is_provisional_serious:
        flow_complete = True
    else:
        # Se for sério, precisa ter respondido até onde o fluxo de correção leva
        if q3_6 in ["Y", "PY"]: flow_complete = True
        elif q3_6 in ["N", "PN", "NI"] and q3_7 in ["Y", "PY"]: flow_complete = True
        elif q3_6 in ["N", "PN", "NI"] and q3_7 in ["N", "PN", "NI"] and q3_8 != "Selecione...": flow_complete = True

if flow_complete:
    # 1. Baseado na combinação inicial (Se não for sério, é o pior entre A e B)
    if not is_provisional_serious:
        if temp_risk_a == "MODERATE" or temp_risk_b == "MODERATE":
            d3_risk = "MODERATE"
            d3_reason = f"Risco Moderado em A ({temp_risk_a}) ou B ({temp_risk_b})."
        else:
            d3_risk = "LOW"
            d3_reason = "Baixo risco de viés de seleção."
            
    # 2. Se entrou no fluxo de correção (Serious)
    else:
        base_reason = f"Viés Sério identificado (A: {temp_risk_a}, B: {temp_risk_b})."
        
        if q3_6 in ["Y", "PY"]:
            d3_risk = "MODERATE"
            d3_reason = base_reason + " Corrigido pela análise (3.6)."
        elif q3_7 in ["Y", "PY"]:
            d3_risk = "MODERATE"
            d3_reason = base_reason + " Mitigado por análise de sensibilidade (3.7)."
        elif q3_8 in ["Y", "PY"]:
            d3_risk = "CRITICAL"
            d3_reason = base_reason + " Viés severo confirmado e não corrigido."
        else:
            d3_risk = "SERIOUS"
            d3_reason = base_reason + " Não corrigido, mas não considerado severo/crítico."

# Salva resultado
risks["D3"] = d3_risk
reasons["D3"] = d3_reason

report_data["domains"]["Domínio 3"] = {
    "risk": d3_risk, 
    "reason": d3_reason, 
    "answers": {"3.1": q3_1, "3.2": q3_2, "3.3": q3_3, "3.4": q3_4, "3.5": q3_5, "3.6": q3_6, "3.7": q3_7, "3.8": q3_8}
}

display_risk_card("Domínio 3", d3_risk, d3_reason)
st.divider()

# --- DOMÍNIO 4: DADOS FALTANTES (LÓGICA CORRIGIDA) ---
st.header("Domínio 4: Viés devido a Dados Faltantes")

st.markdown("""
Este domínio avalia a integridade dos dados e a estratégia de análise.
O algoritmo calcula o risco assim que uma conclusão é atingida (Early Exit).
""")

# --- PASSO 1: TRIAGEM (4.1 a 4.3) ---
c1_d4, c2_d4 = st.columns(2)

with c1_d4:
    q4_1 = st.selectbox(
        "4.1 Dados da intervenção completos?",
        ["Selecione...", "Y", "PY", "PN", "N", "NI"],
        help="Y/PY: Quase todos os participantes têm dados da intervenção."
    )
    q4_3 = st.selectbox(
        "4.3 Dados de confundidores (covariáveis) completos?",
        ["Selecione...", "Y", "PY", "PN", "N", "NI"],
        help="Y/PY: Quase todos os participantes têm dados das variáveis de ajuste."
    )

with c2_d4:
    q4_2 = st.selectbox(
        "4.2 Dados do desfecho completos?",
        ["Selecione...", "Y", "PY", "PN", "N", "NI"],
        help="Y/PY: Quase todos os participantes têm dados do desfecho."
    )

# Verifica se há dados faltantes (Any N/PN/NI)
missing_data = False
if "Selecione..." not in [q4_1, q4_2, q4_3]:
    if q4_1 in ["PN", "N", "NI"] or q4_2 in ["PN", "N", "NI"] or q4_3 in ["PN", "N", "NI"]:
        missing_data = True

# --- PASSO 2: ESTRATÉGIA DE ANÁLISE (4.4) ---
q4_4 = "NA"
analysis_type = "NONE"

if missing_data:
    st.divider()
    st.subheader("Estratégia de Análise")
    
    q4_4 = st.selectbox(
        "4.4 A análise foi feita apenas com casos completos?",
        ["Selecione...", "Y", "PY", "PN", "N", "NI"],
        help="Y/PY/NI: Segue para análise de Casos Completos.\nN/PN: Segue para Imputação."
    )
    
    if q4_4 in ["Y", "PY", "NI"]: analysis_type = "COMPLETE_CASE"
    elif q4_4 in ["N", "PN"]: analysis_type = "IMPUTATION_OR_OTHER"

# --- PASSO 3: RAMIFICAÇÃO E EARLY EXIT ---
q4_5, q4_6 = "NA", "NA"
q4_7, q4_8, q4_9, q4_10 = "NA", "NA", "NA", "NA"
need_4_11 = False

# ==========================================
# RAMO A: CASOS COMPLETOS (4.4 Y/PY/NI)
# ==========================================
if analysis_type == "COMPLETE_CASE":
    st.markdown("**Avaliação: Análise de Casos Completos**")
    
    # 4.5 Sempre aparece neste ramo
    q4_5 = st.selectbox(
        "4.5 A exclusão está relacionada ao valor real do desfecho (MNAR)?",
        ["Selecione...", "Y", "PY", "PN", "N", "NI"],
        help="N/PN: Não relacionado -> Baixo Risco (Early Exit).\nY/PY/NI: Possível viés -> Pergunta 4.6."
    )
    
    # Lógica Corrigida: 4.6 só aparece se 4.5 for RUIM (Y/PY/NI)
    if q4_5 in ["Y", "PY", "NI"]:
        q4_6 = st.selectbox(
            "4.6 A relação entre perda e desfecho é explicada pelo modelo?",
            ["Selecione...", "Y", "PY", "WN", "NI", "SN"],
            help="Verifica se variáveis de ajuste corrigem o viés (MAR)."
        )
        # Se precisou responder 4.6, precisaremos de 4.11 para confirmar
        if q4_6 != "Selecione...": need_4_11 = True

# ==========================================
# RAMO B: IMPUTAÇÃO / OUTROS (4.4 N/PN)
# ==========================================
elif analysis_type == "IMPUTATION_OR_OTHER":
    st.markdown("**Avaliação: Imputação / Outros Métodos**")
    
    q4_7 = st.selectbox(
        "4.7 A análise foi baseada em imputação de valores?",
        ["Selecione...", "Y", "PY", "PN", "N", "NI"]
    )
    
    # Ramo B1: Imputação
    if q4_7 in ["Y", "PY"]:
        q4_8 = st.selectbox(
            "4.8 As premissas MAR/MCAR são razoáveis?",
            ["Selecione...", "Y", "PY", "PN", "N", "NI"]
        )
        
        if q4_8 in ["Y", "PY"]:
            q4_9 = st.selectbox(
                "4.9 A imputação foi apropriada?",
                ["Selecione...", "Y", "PY", "WN", "NI", "SN"],
                help="Y/PY: Leva a Baixo Risco (Early Exit)."
            )
        
        # Gatilhos para 4.11
        if q4_8 in ["N", "PN", "NI"]: need_4_11 = True
        elif q4_9 in ["WN", "NI", "SN"]: need_4_11 = True

    # Ramo B2: Outros Métodos
    elif q4_7 in ["N", "PN", "NI"]:
        q4_10 = st.selectbox(
            "4.10 Foi usado outro método apropriado (ex: IPW)?",
            ["Selecione...", "Y", "PY", "WN", "NI", "SN"],
            help="Y/PY: Leva a Baixo Risco (Early Exit)."
        )
        if q4_10 in ["WN", "NI", "SN"]: need_4_11 = True

# --- PASSO 4: EVIDÊNCIA DE NÃO-VIÉS (4.11) ---
q4_11 = "NA"
if need_4_11:
    st.divider()
    help_4_11 = "Existem evidências (ex: análise de sensibilidade) de que o resultado NÃO é enviesado?"
    q4_11 = st.selectbox(
        "4.11 Existe evidência de que o resultado não é enviesado?",
        ["Selecione...", "Y", "PY", "PN", "N", "NI"],
        help=help_4_11
    )

# --- CÁLCULO DE RISCO (EARLY EXIT) ---
d4_risk = "PENDENTE"
d4_reason = "Aguardando respostas..."

# 1. Triagem Inicial (All Low)
if not missing_data:
    if "Selecione..." not in [q4_1, q4_2, q4_3]:
        d4_risk = "LOW"
        d4_reason = "Dados completos para intervenção, desfecho e confundidores."

# 2. Early Exit: Casos Completos -> Não relacionado ao desfecho (4.5 N/PN)
elif analysis_type == "COMPLETE_CASE" and q4_5 in ["N", "PN"]:
    d4_risk = "LOW"
    d4_reason = "Exclusão de participantes não relacionada ao desfecho (Baixo risco de viés)."

# 3. Early Exit: Sucesso na Imputação
elif q4_9 in ["Y", "PY"]:
    d4_risk = "LOW"
    d4_reason = "Imputação apropriada com premissas válidas."

# 4. Early Exit: Sucesso em Outros Métodos
elif q4_10 in ["Y", "PY"]:
    d4_risk = "LOW"
    d4_reason = "Método alternativo apropriado utilizado."

# 5. Caminhos de Falha / Mitigação (Requer 4.11)
elif need_4_11 and q4_11 != "Selecione...":
    
    came_from_strong_no = (q4_6 == "SN") or (q4_9 == "SN") or (q4_10 == "SN")
    
    if q4_11 in ["Y", "PY"]:
        if came_from_strong_no:
            d4_risk = "SERIOUS" 
            d4_reason = "Erro grave mitigado parcialmente."
        else:
            d4_risk = "MODERATE"
            d4_reason = "Problemas mitigados por análise de sensibilidade."
            
    elif q4_11 in ["N", "PN", "NI"]:
        if came_from_strong_no:
            d4_risk = "CRITICAL"
            d4_reason = "Falha metodológica grave sem evidência de robustez."
        else:
            d4_risk = "SERIOUS"
            d4_reason = "Viés de dados faltantes provável e não mitigado."

# Salva resultado
risks["D4"] = d4_risk
reasons["D4"] = d4_reason

report_data["domains"]["Domínio 4"] = {
    "risk": d4_risk, 
    "reason": d4_reason, 
    "answers": {k: v for k, v in locals().items() if k.startswith('q4_') and isinstance(v, str)}
}

display_risk_card("Domínio 4", d4_risk, d4_reason)
st.divider()

# --- DOMÍNIO 5: MEDIÇÃO ---
st.header("Domínio 5: Medição do Desfecho")
c1, c2 = st.columns(2)
with c1:
    q5_1 = st.selectbox("5.1 Métodos diferem entre grupos?", ["Selecione...", "Y", "PY", "PN", "N", "NI"])
    q5_2 = st.selectbox("5.2 Avaliadores cientes da intervenção?", ["Selecione...", "Y", "PY", "PN", "N", "NI"])
with c2:
    q5_3 = st.selectbox("5.3 Avaliação influenciada pelo conhecimento?", ["Selecione...", "NA", "SY", "WY", "PN", "N", "NI"])

d5_risk, d5_reason = "PENDENTE", "Aguardando respostas..."
if "Selecione..." not in [q5_1, q5_2, q5_3]:
    if q5_1 in ["Y", "PY"]: d5_risk, d5_reason = "SERIOUS", "Métodos de medição diferentes entre os grupos."
    elif q5_2 in ["Y", "PY", "NI"]:
        if q5_3 == "SY": d5_risk, d5_reason = "SERIOUS", "Avaliação subjetiva influenciada pelo conhecimento da intervenção."
        elif q5_3 in ["WY", "NI"]: d5_risk, d5_reason = "MODERATE", "Possível influência no avaliador."
        else: d5_risk, d5_reason = "LOW", "Avaliador ciente, mas desfecho objetivo."
    else:
        if q5_1 == "NI": d5_risk, d5_reason = "MODERATE", "Avaliador cego, mas incerteza sobre comparabilidade dos métodos."
        else: d5_risk, d5_reason = "LOW", "Medição objetiva e comparável."

risks["D5"] = d5_risk
reasons["D5"] = d5_reason
report_data["domains"]["Domínio 5"] = {"risk": d5_risk, "reason": d5_reason, "answers": {"5.1": q5_1, "5.2": q5_2, "5.3": q5_3}}
display_risk_card("Domínio 5", d5_risk, d5_reason)
st.divider()

# --- DOMÍNIO 6: RELATO SELETIVO ---
st.header("Domínio 6: Relato Seletivo")
c1, c2 = st.columns(2)
with c1:
    q6_1 = st.selectbox("6.1 Relatado conforme plano prévio?", ["Selecione...", "Y", "PY", "PN", "N", "NI"])
    q6_2 = st.selectbox("6.2 Seleção baseada em múltiplas medidas?", ["Selecione...", "Y", "PY", "PN", "N", "NI"])
with c2:
    q6_3 = st.selectbox("6.3 Seleção baseada em múltiplas análises?", ["Selecione...", "Y", "PY", "PN", "N", "NI"])
    q6_4 = st.selectbox("6.4 Seleção baseada em subgrupos?", ["Selecione...", "Y", "PY", "PN", "N", "NI"])

d6_risk, d6_reason = "PENDENTE", "Aguardando respostas..."
if "Selecione..." not in [q6_1, q6_2, q6_3, q6_4]:
    if q6_1 in ["Y", "PY"]: d6_risk, d6_reason = "LOW", "Seguiu plano de análise pré-especificado."
    else:
        count_selection = 0
        if q6_2 in ["Y", "PY"]: count_selection += 1
        if q6_3 in ["Y", "PY"]: count_selection += 1
        if q6_4 in ["Y", "PY"]: count_selection += 1
        
        count_ni = 0
        if q6_2 == "NI": count_ni += 1
        if q6_3 == "NI": count_ni += 1
        if q6_4 == "NI": count_ni += 1

        if count_selection >= 2: d6_risk, d6_reason = "CRITICAL", "Fortes evidências de seleção de resultados (P-hacking) em múltiplos aspectos."
        elif count_selection == 1: d6_risk, d6_reason = "SERIOUS", "Evidência de seleção em um aspecto (medida, análise ou subgrupo)."
        elif count_ni == 3: d6_risk, d6_reason = "SERIOUS", "Sem plano de análise e sem informação suficiente para julgar seleção."
        elif count_ni > 0: d6_risk, d6_reason = "MODERATE", "Sem plano de análise e algumas informações faltando."
        else: d6_risk, d6_reason = "MODERATE", "Sem plano de análise, mas sem evidências claras de seleção."

risks["D6"] = d6_risk
reasons["D6"] = d6_reason
report_data["domains"]["Domínio 6"] = {"risk": d6_risk, "reason": d6_reason, "answers": {"6.1": q6_1, "6.2": q6_2, "6.3": q6_3, "6.4": q6_4}}
display_risk_card("Domínio 6", d6_risk, d6_reason)
st.divider()

# --- CÁLCULO GERAL ALGORITMO ---
st.header("Julgamento de Risco (Overall)")
all_risks = list(risks.values())
algo_risk = "PENDENTE"

if "PENDENTE" in all_risks:
    st.warning("Responda todos os domínios para ver o cálculo.")
else:
    # Se Domínio 1 estava em construção (N/A), ignoramos ele no cálculo geral por enquanto
    valid_risks = [r for r in all_risks if r != "N/A"]
    
    if "CRITICAL" in valid_risks: algo_risk = "CRITICAL"
    elif valid_risks.count("SERIOUS") >= 2: algo_risk = "CRITICAL"
    elif "SERIOUS" in valid_risks: algo_risk = "SERIOUS"
    elif valid_risks.count("MODERATE") >= 3: algo_risk = "SERIOUS"
    elif "MODERATE" in valid_risks: algo_risk = "MODERATE"
    else: algo_risk = "LOW"
    
    st.markdown(f"""
    <div style="padding: 15px; background-color: {get_risk_color(algo_risk)}; color: white; text-align: center; border-radius: 8px;">
        <h3>RISCO SUGERIDO (ALGORITMO): {algo_risk}</h3>
    </div>
    """, unsafe_allow_html=True)

# --- JULGAMENTO DO PESQUISADOR ---
st.markdown("### Validação pelo Pesquisador")
st.caption("O algoritmo oferece uma sugestão padrão. O pesquisador pode alterar o julgamento final se houver justificativa (Guidance Note 17).")

col_final1, col_final2 = st.columns([1, 2])
with col_final1:
    manual_risk = st.selectbox(
        "Decisão Final de Risco Global",
        ["LOW", "MODERATE", "SERIOUS", "CRITICAL"],
        index=["LOW", "MODERATE", "SERIOUS", "CRITICAL"].index(algo_risk) if algo_risk != "PENDENTE" else 0
    )
with col_final2:
    manual_justification = st.text_area(
        "Justificativa do Pesquisador (Obrigatório para Override)",
        placeholder="Explique se concordou com o algoritmo ou por que alterou o risco..."
    )

# --- ÁREA DE DOWNLOAD ---
st.divider()
st.subheader("📄 Exportar Relatório")

if st.button("Gerar Arquivos para Download"):
    # Atualiza dados finais
    report_data["algo_risk"] = algo_risk
    report_data["manual_risk"] = manual_risk
    report_data["manual_justification"] = manual_justification
    
    # Gera arquivos
    try:
        docx_file = generate_docx(report_data)
        pdf_file = generate_pdf(report_data)
        
        col_d1, col_d2 = st.columns(2)
        
        with col_d1:
            st.download_button(
                label="📥 Baixar Relatório WORD (.docx)",
                data=docx_file.getvalue(),
                file_name=f"ROBINS_I_{study_id}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        
        with col_d2:
            st.download_button(
                label="📥 Baixar Relatório PDF (.pdf)",
                data=pdf_file,
                file_name=f"ROBINS_I_{study_id}.pdf",
                mime="application/pdf"
            )
    except Exception as e:
        st.error(f"Erro ao gerar arquivos: {e}")